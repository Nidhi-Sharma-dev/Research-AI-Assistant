"""
Export utilities – generate PDF and DOCX grant proposals from the report.
"""
import os
import re
import logging
from datetime import datetime
from typing import Dict, Any

logger = logging.getLogger(__name__)


def export_proposal_pdf(grant: Dict[str, Any], output_dir: str = "./outputs") -> str:
    """Generate a PDF of the grant proposal. Returns path."""
    try:
        from fpdf import FPDF

        os.makedirs(output_dir, exist_ok=True)
        fname = f"grant_proposal_{grant.get('agency','NSF')}_{datetime.now():%Y%m%d_%H%M%S}.pdf"
        fpath = os.path.join(output_dir, fname)

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Title page
        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 12, "GRANT PROPOSAL", ln=True, align="C")
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, f"Agency: {grant.get('agency', '')}", ln=True, align="C")
        pdf.set_font("Helvetica", "", 12)
        pdf.cell(0, 8, f"Title: {grant.get('research_topic', '')}", ln=True, align="C")
        pdf.cell(0, 8, f"PI: {grant.get('pi_name', '')} | {grant.get('institution', '')}", ln=True, align="C")
        pdf.cell(0, 8, f"Budget: {grant.get('budget_total','')} | Duration: {grant.get('duration_years','')} years", ln=True, align="C")
        pdf.ln(10)

        for section, content in grant.get("sections", {}).items():
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_fill_color(220, 230, 255)
            pdf.cell(0, 9, section, ln=True, fill=True)
            pdf.ln(2)
            pdf.set_font("Helvetica", "", 11)
            safe_content = content.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(0, 6, safe_content)
            pdf.ln(5)

        pdf.output(fpath)
        logger.info("PDF saved: %s", fpath)
        return fpath

    except Exception as exc:
        logger.error("PDF export failed: %s", exc)
        return ""


def export_proposal_docx(grant: Dict[str, Any], output_dir: str = "./outputs") -> str:
    """Generate a DOCX of the grant proposal. Returns path."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        os.makedirs(output_dir, exist_ok=True)
        fname = f"grant_proposal_{grant.get('agency','NSF')}_{datetime.now():%Y%m%d_%H%M%S}.docx"
        fpath = os.path.join(output_dir, fname)

        doc = Document()
        doc.core_properties.title   = grant.get("research_topic", "Grant Proposal")
        doc.core_properties.author  = grant.get("pi_name", "")
        doc.core_properties.subject = f"{grant.get('agency','')} Grant Proposal"

        # Title
        title = doc.add_heading(grant.get("research_topic", "Grant Proposal"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(
            f"Agency: {grant.get('agency','')}  |  PI: {grant.get('pi_name','')}  |  "
            f"Institution: {grant.get('institution','')}  |  "
            f"Budget: {grant.get('budget_total','')}  |  "
            f"Duration: {grant.get('duration_years','')} years"
        ).italic = True

        doc.add_paragraph()

        for section, content in grant.get("sections", {}).items():
            doc.add_heading(section, level=1)
            doc.add_paragraph(content)
            doc.add_paragraph()

        doc.save(fpath)
        logger.info("DOCX saved: %s", fpath)
        return fpath

    except Exception as exc:
        logger.error("DOCX export failed: %s", exc)
        return ""


def export_report_markdown(report_dict: Dict[str, Any], output_dir: str = "./outputs") -> str:
    """Export the full research report as markdown."""
    os.makedirs(output_dir, exist_ok=True)
    fname = f"research_report_{datetime.now():%Y%m%d_%H%M%S}.md"
    fpath = os.path.join(output_dir, fname)

    lines = ["# AI Research Assistant – Full Report\n"]
    lines.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n---\n")

    for section, data in report_dict.items():
        lines.append(f"## {section.title()}\n")
        if isinstance(data, dict):
            for k, v in data.items():
                if isinstance(v, list):
                    lines.append(f"**{k}:**")
                    for item in v:
                        lines.append(f"  - {item}")
                    lines.append("")
                elif isinstance(v, str) and len(v) > 200:
                    lines.append(f"**{k}:**\n\n{v}\n")
                else:
                    lines.append(f"**{k}:** {v}")
        lines.append("\n---\n")

    with open(fpath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return fpath
